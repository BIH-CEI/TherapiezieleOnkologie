"""Fügt ein Deckblatt als erste Seite in das gerenderte Master-DOCX ein.

Wird am Ende von build_master.py aufgerufen. Arbeitet auf der bereits
gerenderten Datei: öffnet sie mit python-docx, fügt Cover-Elemente vor
dem ersten Inhaltselement ein, setzt Seitenumbruch, speichert zurück.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).parent
LOGO = HERE / "bih-logo.png"

BIH_NAVY = RGBColor(0x00, 0x37, 0x54)           # dk1 — Titel
BIH_DARKRED = RGBColor(0xAF, 0x18, 0x21)         # dk2/accent1 — Primär-Akzent
BIH_CORAL = RGBColor(0xEA, 0x54, 0x51)           # accent2 — Cover-Label, Sekundär
BIH_TEAL = RGBColor(0x00, 0x9A, 0xA9)            # accent5 — Trennlinie, dezent
BIH_GREY = RGBColor(0x9D, 0x9D, 0x9D)
# Alias-Namen (rückwärts-kompatibel mit der bisherigen Cover-Logik)
BIH_BLUE = BIH_DARKRED        # Subtitle in Dunkelrot statt Teal
BIH_MAGENTA = BIH_CORAL       # Cover-Label in Koral
BIH_SLATE = BIH_NAVY          # Metadaten-Werte in Navy


def _set_font(run, name="Fira Sans", size=None, bold=None, color=None):
    run.font.name = name
    rpr = run.font.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_horizontal_rule(paragraph, color="70ACC0", width_pt=6):
    """Setzt eine farbige Linie unter dem Absatz."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(width_pt * 4))  # half-points scaled
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _insert_paragraph_before(reference_para, text="", style=None):
    """Fügt einen neuen Absatz VOR dem Referenz-Absatz ein."""
    new_p = OxmlElement("w:p")
    reference_para._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, reference_para._parent)
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p


def _make_paragraph_before_xml(reference_element, body):
    """Erzeugt eine python-docx Paragraph, deren XML vor reference_element steht."""
    new_p = OxmlElement("w:p")
    reference_element.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    # Parent muss ein Container sein (kein _Body); python-docx erlaubt None für add_run
    return Paragraph(new_p, body)


def inject_cover(docx_path: Path, *, title: str, subtitle: str, metadata: dict[str, str]) -> None:
    doc = Document(docx_path)
    body = doc.element.body
    # Erstes Body-Kind-Element finden — egal ob <w:p>, <w:sdt>, <w:tbl>, …
    # (NICHT doc.paragraphs[0], denn das überspringt SDTs/Tables und liefert
    # erst die erste Body-Paragraph zurück, was bei TOC-SDT zu Insertion nach
    # der TOC führen würde.)
    first_body_child = None
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag in ("p", "sdt", "tbl"):
            first_body_child = child
            break
    if first_body_child is None:
        raise RuntimeError("Konnte kein erstes Body-Element finden")

    # Anker für Inserts: ein leerer Paragraph direkt vor dem ersten Body-Kind,
    # in den wir den abschließenden Seitenumbruch packen. Anker bleibt am ENDE
    # der Cover-Seite (kurz vor TOC). Mit addprevious produziert die
    # Insert-REIHENFOLGE die visuelle Top-down-Reihenfolge.
    pb_para = _make_paragraph_before_xml(first_body_child, body)
    pb_para.add_run().add_break(WD_BREAK.PAGE)
    anchor = pb_para

    # Top-Padding
    for _ in range(2):
        _insert_paragraph_before(anchor)

    # Kontext-Label oben
    lbl_p = _insert_paragraph_before(anchor)
    lbl_run = lbl_p.add_run("BIH-CEI · Berlin Institute of Health at Charité")
    _set_font(lbl_run, size=10, bold=True, color=BIH_MAGENTA)

    # Großer Abstand zum Titel
    for _ in range(4):
        _insert_paragraph_before(anchor)

    # Titel
    title_p = _insert_paragraph_before(anchor)
    title_run = title_p.add_run(title)
    _set_font(title_run, size=32, bold=True, color=BIH_NAVY)

    # Subtitle
    sub_p = _insert_paragraph_before(anchor)
    sub_run = sub_p.add_run(subtitle)
    _set_font(sub_run, size=16, bold=False, color=BIH_BLUE)

    # Vertikaler Abstand
    for _ in range(2):
        _insert_paragraph_before(anchor)

    # Horizontale Linie
    rule_p = _insert_paragraph_before(anchor)
    _add_horizontal_rule(rule_p, color="AF1821", width_pt=4)

    # Abstand
    _insert_paragraph_before(anchor)

    # Metadaten-Block (Liefergegenstand, Auftrag, Stand, Projektleitung, FHIR-Lead)
    for key, value in metadata.items():
        p = _insert_paragraph_before(anchor)
        r_key = p.add_run(f"{key}:  ")
        _set_font(r_key, size=11, bold=True, color=BIH_NAVY)
        r_val = p.add_run(value)
        _set_font(r_val, size=11, color=BIH_SLATE)

    # Bottom-Padding
    for _ in range(4):
        _insert_paragraph_before(anchor)

    doc.save(docx_path)


def main():
    import sys
    if len(sys.argv) < 2:
        print("Usage: inject_cover.py <docx>")
        sys.exit(1)
    docx_path = Path(sys.argv[1])
    inject_cover(
        docx_path,
        title="LG-01 Analysebericht",
        subtitle='Spezifikation „Onkologische Therapieziele"',
        metadata={
            "Liefergegenstand": "LG-01 · Fachliche Analyse",
            "Auftrag": "Gematik C002717",
            "Stand": date.today().strftime("%d.%m.%Y") + " · Entwurf",
            "Projektleitung": "Prof. Dr. Sylvia Thun",
            "Technische Leitung FHIR": "Thomas Debertshäuser",
        },
    )
    print(f"Cover injected in: {docx_path}")


if __name__ == "__main__":
    main()
