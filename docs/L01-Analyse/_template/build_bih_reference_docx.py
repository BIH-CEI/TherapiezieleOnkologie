"""
Erzeugt eine unoffizielle BIH-Reference-DOCX zur Verwendung mit pandoc --reference-doc.

Quelle der Brand-Assets: BIH PowerPoint-Templates (bih-presentation Skill),
extrahiert aus BIH_16x9_latest.potx.

WICHTIG: Dies ist KEIN offizielles BIH-Word-Template. Es lehnt sich am
Corporate Design an (Fira Sans, BIH-Blau #70ACC0, BIH-Magenta #C12075,
Dunkelnavy #003754), ersetzt aber keine offizielle Charité/BIH-Vorlage
fuer extern verteilte Dokumente.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Mm, Pt, RGBColor

HERE = Path(__file__).parent
OUTPUT = HERE / "BIH_Report_unofficial_reference.docx"
LOGO = HERE / "bih-logo.png"

# BIH-Farben (aus BIH_16x9.pptx vom März 2026 — neues CI)
BIH_NAVY = RGBColor(0x00, 0x37, 0x54)      # Primär
BIH_TEAL = RGBColor(0x00, 0x9A, 0xA9)      # Sekundär (löst altes #70ACC0 ab)
BIH_CORAL = RGBColor(0xEA, 0x54, 0x51)     # Akzent (löst altes Magenta ab)
BIH_DARKRED = RGBColor(0xAF, 0x18, 0x21)   # Akzent dunkel
BIH_PURPLE = RGBColor(0x78, 0x76, 0xB6)    # Tertiär
BIH_GOLD = RGBColor(0x9D, 0x72, 0x20)
BIH_LIGHT_CORAL = RGBColor(0xFF, 0xB0, 0xAC)
BIH_GREY = RGBColor(0x9D, 0x9D, 0x9D)
BLACK = RGBColor(0x00, 0x00, 0x00)

# Aliase nach Theme-Rollen (theme1.xml: dk1/dk2/accent1-6)
# H-Hierarchie: H1 Navy → H2 Dunkelrot (dk2) → H3 Teal (accent5) → H4 Purple (accent6)
BIH_BLUE = BIH_DARKRED      # H2 Primär-Akzent (#AF1821)
BIH_MAGENTA = BIH_CORAL     # Cover-Label und Sekundärakzent (#EA5451)
BIH_PLUM = BIH_TEAL         # H3 Tertiär-Akzent (#009AA9)
BIH_SLATE = BIH_PURPLE      # H4 (#7876B6)
BIH_TAN = BIH_GOLD

FONT_BODY = "Fira Sans"
FONT_HEAD = "Fira Sans"


def set_font(run_or_style, font_name=FONT_BODY, size=None, bold=None, color=None):
    """Setzt Schriftattribute robust auf Run oder Style."""
    font = run_or_style.font
    font.name = font_name
    # East-Asian + complex-script auch setzen, sonst weicht Word aus
    rpr = font.element.get_or_add_rPr()
    for tag in ("w:rFonts",):
        existing = rpr.find(qn(tag))
        if existing is None:
            existing = OxmlElement(tag)
            rpr.append(existing)
        existing.set(qn("w:ascii"), font_name)
        existing.set(qn("w:hAnsi"), font_name)
        existing.set(qn("w:cs"), font_name)
        existing.set(qn("w:eastAsia"), font_name)
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if color is not None:
        font.color.rgb = color


def _set_body_spacing(style):
    """Einheitliche Body-Spacing-Werte für alle pandoc-relevanten Body-Styles."""
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25


def configure_styles(doc):
    styles = doc.styles

    # Body / Normal — KEINE explizite Farbe (sonst überschreibt sie
    # tblStylePr-Header-Farben in Tabellen). Word-Default = automatic (schwarz).
    normal = styles["Normal"]
    set_font(normal, FONT_BODY, size=11)
    _set_body_spacing(normal)

    # Pandoc-typische Body-Varianten alle auf gleiche Spacing bringen.
    style_names_in_doc = {s.name for s in styles}
    for name in ("Body Text", "BodyText", "First Paragraph", "FirstParagraph", "Compact"):
        if name in style_names_in_doc:
            s = styles[name]
            set_font(s, FONT_BODY, size=11)
            _set_body_spacing(s)
            # Compact darf in Tabellen tighter sein — wird vom tablestyle ggf. überschrieben

    # Title (Berichtstitel)
    title = styles["Title"]
    set_font(title, FONT_HEAD, size=28, bold=True, color=BIH_NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)

    # Subtitle
    if "Subtitle" in [s.name for s in styles]:
        subtitle = styles["Subtitle"]
        set_font(subtitle, FONT_HEAD, size=14, bold=False, color=BIH_BLUE)
        subtitle.paragraph_format.space_after = Pt(18)

    # Headings 1-4
    for level, size, color, bold in [
        (1, 18, BIH_NAVY, True),
        (2, 14, BIH_BLUE, True),
        (3, 12, BIH_PLUM, True),
        (4, 11, BIH_SLATE, True),
    ]:
        style_name = f"Heading {level}"
        h = styles[style_name]
        set_font(h, FONT_HEAD, size=size, bold=bold, color=color)
        h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True

    # Quote / Intense Quote (für Hervorhebungen)
    for q_name in ("Quote", "Intense Quote"):
        if q_name in [s.name for s in styles]:
            qstyle = styles[q_name]
            set_font(qstyle, FONT_BODY, size=11, color=BIH_PLUM)
            qstyle.paragraph_format.left_indent = Cm(0.5)

    # Caption (Bild-/Tabellenunterschrift)
    if "Caption" in [s.name for s in styles]:
        cap = styles["Caption"]
        set_font(cap, FONT_BODY, size=9, color=BIH_SLATE)

    # Liste
    for list_name in ("List Bullet", "List Number"):
        if list_name in [s.name for s in styles]:
            ls = styles[list_name]
            set_font(ls, FONT_BODY, size=11, color=BLACK)


def configure_table_grid_style(doc):
    """Setzt Standard-Tabellenstil auf BIH-Blau-Header."""
    # Pandoc nutzt typischerweise 'Table Grid' für markdown-Tabellen.
    styles = doc.styles
    if "Table Grid" not in [s.name for s in styles]:
        return
    tg = styles["Table Grid"]
    # Set Tabelle-Font
    set_font(tg, FONT_BODY, size=10)


def add_page_setup(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)


def add_header_with_logo(doc):
    section = doc.sections[0]
    header = section.header
    # Lösche bestehende paragraph(en) im Header — wir wollen frisch
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run()
    if LOGO.exists():
        run.add_picture(str(LOGO), width=Cm(3.5))


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("BIH-CEI · Gematik KIG Onkologische Therapieziele")
    set_font(run, FONT_BODY, size=8, color=BIH_GREY)
    # Seitenzahl rechts (separate Tabstop-Lösung wäre sauberer; hier minimal)


def add_table_style_to_styles_xml(docx_path):
    """Patcht styles.xml im fertigen docx, sodass der Style "Table"
    (von pandoc für Markdown-Tabellen verwendet) mit BIH-Borders und
    Header-Shading existiert.
    """
    import shutil
    import tempfile
    import zipfile

    with tempfile.TemporaryDirectory() as tmp:
        with zipfile.ZipFile(docx_path, "r") as z:
            z.extractall(tmp)
        styles_path = Path(tmp) / "word" / "styles.xml"
        text = styles_path.read_text(encoding="utf-8")

        # Style "Compact" einfügen (pandoc benutzt diesen in Tabellenzellen
        # und tight lists; fehlt er, rendern manche Engines die Zelltexte ausserhalb).
        if 'w:styleId="Compact"' not in text:
            compact_style = (
                '<w:style w:type="paragraph" w:styleId="Compact">'
                '<w:name w:val="Compact"/>'
                '<w:basedOn w:val="Normal"/>'
                '<w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
                '<w:rPr><w:rFonts w:ascii="Fira Sans" w:hAnsi="Fira Sans" w:cs="Fira Sans"/><w:sz w:val="20"/></w:rPr>'
                '</w:style>'
            )
            text = text.replace("</w:styles>", compact_style + "</w:styles>")

        # Style "Table" einfügen falls fehlend.
        if 'w:styleId="Table"' not in text:
            table_style = (
                '<w:style w:type="table" w:styleId="Table">'
                '<w:name w:val="Table" />'
                '<w:basedOn w:val="TableNormal" />'
                '<w:uiPriority w:val="59" />'
                '<w:rPr><w:rFonts w:ascii="Fira Sans" w:hAnsi="Fira Sans" w:cs="Fira Sans"/><w:sz w:val="20"/></w:rPr>'
                '<w:tblPr>'
                '<w:tblBorders>'
                '<w:top w:val="single" w:sz="4" w:space="0" w:color="AF1821"/>'
                '<w:left w:val="single" w:sz="4" w:space="0" w:color="AF1821"/>'
                '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="AF1821"/>'
                '<w:right w:val="single" w:sz="4" w:space="0" w:color="AF1821"/>'
                '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="AF1821"/>'
                '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="AF1821"/>'
                '</w:tblBorders>'
                '</w:tblPr>'
                '<w:tblStylePr w:type="firstRow">'
                '<w:rPr><w:b/><w:color w:val="FFFFFF"/></w:rPr>'
                '<w:tcPr><w:shd w:val="clear" w:color="auto" w:fill="003754"/></w:tcPr>'
                '</w:tblStylePr>'
                '</w:style>'
            )
            text = text.replace("</w:styles>", table_style + "</w:styles>")
        styles_path.write_text(text, encoding="utf-8")

        # Re-zip
        out_tmp = Path(tmp) / "out.docx"
        with zipfile.ZipFile(out_tmp, "w", zipfile.ZIP_DEFLATED) as z:
            for f in Path(tmp).rglob("*"):
                if f.is_file() and f.name != "out.docx":
                    z.write(f, f.relative_to(tmp))
        shutil.copy(out_tmp, docx_path)


def main():
    doc = Document()
    configure_styles(doc)
    configure_table_grid_style(doc)
    add_page_setup(doc)
    add_header_with_logo(doc)
    add_footer(doc)

    # Minimaler Beispielinhalt, damit pandoc beim Übernehmen aller Style-Definitionen
    # sicher den vollen Satz an Styles im Reference-Dokument findet.
    doc.add_paragraph("Berichtstitel", style="Title")
    doc.add_paragraph("Untertitel — Liefergegenstand und Datum", style="Subtitle")
    doc.add_paragraph("Überschrift 1", style="Heading 1")
    doc.add_paragraph("Überschrift 2", style="Heading 2")
    doc.add_paragraph("Überschrift 3", style="Heading 3")
    doc.add_paragraph("Überschrift 4", style="Heading 4")
    doc.add_paragraph(
        "Dies ist ein Absatz im Body-Stil. Schrift Fira Sans, 11 pt, Zeilenabstand 1,25, "
        "BIH-Corporate-Farben in den Überschriften."
    )
    # Liste
    doc.add_paragraph("Ein Bullet-Punkt", style="List Bullet")
    doc.add_paragraph("Ein nummerierter Punkt", style="List Number")
    # Quote
    if "Quote" in [s.name for s in doc.styles]:
        doc.add_paragraph("Hervorgehobener Hinweis als Zitat.", style="Quote")

    # Beispiel-Tabelle, damit pandoc den Tabellenstil findet und übernimmt.
    # Pandoc setzt für Markdown-Tabellen den Style "Table" (mit Border-Layout).
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Spalte A"
    table.rows[0].cells[1].text = "Spalte B"
    table.rows[1].cells[0].text = "Wert 1"
    table.rows[1].cells[1].text = "Wert 2"
    # Header-Zeile blau hinterlegen (BIH-Blau hell)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for cell in table.rows[0].cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "70ACC0")
        cell._tc.get_or_add_tcPr().append(shading)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.save(OUTPUT)
    add_table_style_to_styles_xml(OUTPUT)
    print(f"Geschrieben: {OUTPUT}")


if __name__ == "__main__":
    main()
